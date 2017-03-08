from docx import Document
import os


class NotesDocument:

    def __init__(self, doc_name="MyNotes", path=""):
        self.__doc_name = doc_name
        self.__path = path
        self.__document = Document()
        self.__document.add_heading('Notes', 0)

    def add_paragraph(self, notes):
        for note in notes:
            text = note.get_date() + " - " + note.get_description()
            for tag in note.get_tags():
                text += '[' + tag + ']'
            self.__document.add_paragraph(text, style='ListBullet')

    def save(self,):
        if not self.__path:
            self.__document.save(os.path.dirname(os.path.abspath(__file__)) + "/" + self.__doc_name + ".docx")
        else:
            self.__document.save(os.path.dirname(str(self.__path) + "/" + self.__doc_name + ".docx"))
